from docx import Document
from typing import Dict
import os


def generate_docx(
    summaries: Dict[str, str],
    output_path: str,
    title: str = "Podcast Summary"
):
    """
    Generate a DOCX with topic-wise summaries.
    """

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    document = Document()
    document.add_heading(title, level=0)

    for topic, summary in summaries.items():
        document.add_heading(topic, level=1)
        document.add_paragraph(summary)

    document.save(output_path)